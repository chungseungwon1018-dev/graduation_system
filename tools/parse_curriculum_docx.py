import re
import json
import os
from docx import Document

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
DOCX_FILES = [
    os.path.join(PROJECT_ROOT, '2021 경영정보학과 교육과정(홈페이지).docx'),
    os.path.join(PROJECT_ROOT, '2022 경영정보학과 교육과정 (홈페이지).docx'),
    os.path.join(PROJECT_ROOT, '2023 경영정보학과 교육과정 (홈페이지).docx')
]

NUM_REGEX = re.compile(r"(\d+(?:[\.,]\d+)?)\s*학점")
YEAR_REGEX = re.compile(r"(202\d)")


def extract_from_paragraphs(doc):
    hits = []
    lib_credits = None
    major_credits = None
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        # 교양 관련
        if '교양' in text and '학점' in text:
            nums = NUM_REGEX.findall(text)
            if nums:
                hits.append({'type': '교양_par', 'text': text, 'numbers': nums})
                # Heuristic: first number encountered for 교양
                if lib_credits is None:
                    lib_credits = float(nums[0].replace(',', '.'))
        # 전공 관련
        if '전공' in text and '학점' in text:
            nums = NUM_REGEX.findall(text)
            if nums:
                hits.append({'type': '전공_par', 'text': text, 'numbers': nums})
                if major_credits is None:
                    major_credits = float(nums[0].replace(',', '.'))
    return lib_credits, major_credits, hits


def extract_from_tables(doc):
    # Look for tables that mention 전공필수/전필 or 전공선택/전선
    found = {'전필': None, '전선': None}
    table_summaries = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for r in table.rows:
            row_texts = [c.text.strip() for c in r.cells]
            rows.append(row_texts)
        joined = '\n'.join([' | '.join(r) for r in rows])
        table_summaries.append({'index': ti, 'rows': rows})

        # search for keywords in table text
        if re.search(r'전공필수|전필', joined):
            # try to find credits numbers in the table
            nums = NUM_REGEX.findall(joined)
            credit = float(nums[0].replace(',', '.')) if nums else None
            found['전필'] = {'table_index': ti, 'credit': credit, 'rows': rows}
        if re.search(r'전공선택|전선', joined):
            nums = NUM_REGEX.findall(joined)
            credit = float(nums[0].replace(',', '.')) if nums else None
            found['전선'] = {'table_index': ti, 'credit': credit, 'rows': rows}
    return found, table_summaries


def parse_file(path):
    result = {
        'file': os.path.basename(path),
        'liberal_arts_credits': None,
        'major_credits': None,
        'major_requirements': {'전필': None, '전선': None},
        'paragraph_hits': [],
        'tables_scanned': 0
    }

    try:
        doc = Document(path)
    except Exception as e:
        result['error'] = str(e)
        return result

    lib_credits, major_credits, hits = extract_from_paragraphs(doc)
    result['liberal_arts_credits'] = lib_credits
    result['major_credits'] = major_credits
    result['paragraph_hits'] = hits

    found, table_summaries = extract_from_tables(doc)
    result['tables_scanned'] = len(table_summaries)
    if found.get('전필'):
        result['major_requirements']['전필'] = found['전필']
    if found.get('전선'):
        result['major_requirements']['전선'] = found['전선']

    return result


if __name__ == '__main__':
    results = []
    for f in DOCX_FILES:
        if os.path.exists(f):
            print(f"Parsing: {os.path.basename(f)}")
            r = parse_file(f)
            results.append(r)
        else:
            print(f"File not found: {f}")

    out_path = os.path.join(PROJECT_ROOT, 'tools', 'curriculum_summary.json')
    with open(out_path, 'w', encoding='utf-8') as wf:
        json.dump(results, wf, ensure_ascii=False, indent=2)

    print(f"Saved summary to {out_path}")
